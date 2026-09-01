"""
word/parse_word.py
Usage:  python parse_word.py <filename>
Place file in input/ folder. Output goes to output/ with a datetime stamp.
"""

import sys, os
from datetime import datetime

def main():
    if len(sys.argv) < 2:
        print("Usage: python parse_word.py <filename>")
        print("Example: python parse_word.py report.docx")
        sys.exit(1)

    filename   = sys.argv[1].strip()
    script_dir = os.path.dirname(os.path.abspath(__file__))
    input_path = os.path.join(script_dir, "input", filename)

    if not os.path.isfile(input_path):
        print(f"ERROR: File not found: {input_path}")
        print(f"Place '{filename}' inside the input/ folder and try again.")
        sys.exit(1)

    # ── install dependency if missing ──────────────────────────────────────
    try:
        from docx import Document
    except ImportError:
        print("python-docx not found. Installing...")
        os.system("pip install python-docx --break-system-packages -q")
        from docx import Document

    lines = []
    lines.append(f"FILE: {filename}")
    lines.append(f"PARSED AT: {datetime.now()}")

    # ── open file ──────────────────────────────────────────────────────────
    try:
        doc = Document(input_path)
    except Exception as e:
        print(f"ERROR opening file: {e}")
        sys.exit(1)

    # ── core properties ────────────────────────────────────────────────────
    lines.append("\n" + "="*60)
    lines.append("CORE PROPERTIES")
    lines.append("="*60)
    try:
        cp = doc.core_properties
        for attr in ["title","author","subject","keywords","description",
                     "category","last_modified_by","created","modified","revision"]:
            try:
                val = getattr(cp, attr, None)
                lines.append(f"{attr}: {val}")
            except Exception as e:
                lines.append(f"{attr}: [ERROR: {e}]")
    except Exception as e:
        lines.append(f"[Could not read core properties: {e}]")

    # ── paragraphs ─────────────────────────────────────────────────────────
    lines.append("\n" + "="*60)
    lines.append("PARAGRAPHS")
    lines.append("="*60)
    try:
        for para_idx, para in enumerate(doc.paragraphs):
            try:
                style = para.style.name if para.style else "Unknown"
            except Exception:
                style = "Unknown"
            try:
                text = para.text  # full paragraph text, no filtering
            except Exception as e:
                text = f"[ERROR: {e}]"
            lines.append(f"[{para_idx}] style={style} | {text}")

            # also dump individual runs so nothing is missed
            try:
                for run_idx, run in enumerate(para.runs):
                    try:
                        lines.append(f"    run[{run_idx}]: {run.text}")
                    except Exception as e:
                        lines.append(f"    run[{run_idx}]: [ERROR: {e}]")
            except Exception as e:
                lines.append(f"    runs: [ERROR: {e}]")
    except Exception as e:
        lines.append(f"[Could not iterate paragraphs: {e}]")

    # ── tables ─────────────────────────────────────────────────────────────
    lines.append("\n" + "="*60)
    lines.append("TABLES")
    lines.append("="*60)
    try:
        for tbl_idx, table in enumerate(doc.tables):
            lines.append(f"\n-- Table {tbl_idx + 1} --")
            try:
                for row_idx, row in enumerate(table.rows):
                    row_cells = []
                    for cell in row.cells:
                        try:
                            row_cells.append(cell.text)
                        except Exception as e:
                            row_cells.append(f"[ERROR: {e}]")
                    lines.append(f"  row[{row_idx}]: {' | '.join(row_cells)}")
            except Exception as e:
                lines.append(f"  [Could not iterate rows: {e}]")
    except Exception as e:
        lines.append(f"[Could not iterate tables: {e}]")

    # ── inline shapes / images (metadata only) ─────────────────────────────
    lines.append("\n" + "="*60)
    lines.append("INLINE SHAPES")
    lines.append("="*60)
    try:
        for shape_idx, shape in enumerate(doc.inline_shapes):
            try:
                lines.append(f"[{shape_idx}] type={shape.type}  width={shape.width}  height={shape.height}")
            except Exception as e:
                lines.append(f"[{shape_idx}] [ERROR: {e}]")
    except Exception as e:
        lines.append(f"[Could not iterate inline shapes: {e}]")

    # ── sections ───────────────────────────────────────────────────────────
    lines.append("\n" + "="*60)
    lines.append("SECTIONS")
    lines.append("="*60)
    try:
        for sec_idx, section in enumerate(doc.sections):
            try:
                lines.append(
                    f"[{sec_idx}] orientation={section.orientation} "
                    f"page_width={section.page_width} page_height={section.page_height}"
                )
            except Exception as e:
                lines.append(f"[{sec_idx}] [ERROR: {e}]")
    except Exception as e:
        lines.append(f"[Could not iterate sections: {e}]")

    # ── write output ───────────────────────────────────────────────────────
    timestamp  = datetime.now().strftime("%Y%m%d_%H%M%S")
    base       = os.path.splitext(filename)[0]
    output_dir = os.path.join(script_dir, "output")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"{base}__{timestamp}.txt")

    try:
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        print(f"Done! Output → {output_path}")
    except Exception as e:
        print(f"ERROR writing output: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()
